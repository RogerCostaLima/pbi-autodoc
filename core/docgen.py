"""
Gera documento Word (.docx) com a documentação técnica do modelo Power BI.
"""

import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .extractor import PBIXModel


# ── Cores temáticas ──────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)   # azul escuro
COLOR_ACCENT = RGBColor(0xF0, 0xA5, 0x00)    # âmbar/gold
COLOR_LIGHT = RGBColor(0xF8, 0xF7, 0xF4)     # off-white
COLOR_BORDER = RGBColor(0xDD, 0xDD, 0xD0)    # cinza suave
COLOR_SUCCESS = RGBColor(0x22, 0xA0, 0x6B)   # verde
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x2E)      # texto principal


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, border_color: str = "DDDDDO"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "DDDDCE")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_header_row(table, headers: list[str], bg: str = "1A1A2E", fg: RGBColor = None):
    if fg is None:
        fg = RGBColor(0xFF, 0xFF, 0xFF)
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        _set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = fg
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_data_row(row, even: bool = True):
    bg = "F8F7F4" if even else "FFFFFF"
    for cell in row.cells:
        _set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = COLOR_TEXT


def _add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(11)
    return p


def _add_kpi_table(doc: Document, model: PBIXModel):
    summary = model.summary
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    labels = ["Tabelas", "Colunas", "Medidas", "Relacionamentos", "Fontes", "Páginas"]
    values = [
        summary["tables"],
        summary["columns"],
        summary["measures"],
        summary["relationships"],
        summary["data_sources"],
        summary["report_pages"],
    ]

    _add_header_row(table, labels, bg="F0A500", fg=RGBColor(0x1A, 0x1A, 0x2E))

    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        _set_cell_bg(cell, "FFFFFF")
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()


def generate_docx(
    model: PBIXModel,
    ai_overview: str = "",
    logo_path: str = None,
    output_path: str = None,
) -> bytes:
    """Gera e retorna o .docx como bytes."""

    doc = Document()

    # ── Margens ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)

    # ── Capa ───────────────────────────────────────────────────────────────────
    if logo_path and Path(logo_path).exists():
        doc.add_picture(logo_path, width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run("Documentação Técnica · Power BI")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = COLOR_PRIMARY

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(model.file_name)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    # ── Resumo executivo ───────────────────────────────────────────────────────
    _add_section_heading(doc, "1. Resumo do Modelo")
    _add_kpi_table(doc, model)

    # ── Visão geral por IA ─────────────────────────────────────────────────────
    if ai_overview:
        _add_section_heading(doc, "2. Visão Geral (IA)")
        p = doc.add_paragraph(ai_overview)
        p.style.font.size = Pt(10)
        doc.add_paragraph()

    section_offset = 3 if ai_overview else 2

    # ── Fontes de dados ────────────────────────────────────────────────────────
    _add_section_heading(doc, f"{section_offset}. Fontes de Dados")
    section_offset += 1

    if model.data_sources:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _add_header_row(table, ["Nome", "Tipo"])
        for i, ds in enumerate(model.data_sources):
            row = table.add_row()
            row.cells[0].text = ds.name
            row.cells[1].text = ds.kind
            _style_data_row(row, even=(i % 2 == 0))
    else:
        doc.add_paragraph("Nenhuma fonte identificada.", style="Body Text")

    doc.add_paragraph()

    # ── Tabelas e Colunas ─────────────────────────────────────────────────────
    _add_section_heading(doc, f"{section_offset}. Tabelas e Colunas")
    section_offset += 1

    for t_idx, table_obj in enumerate(model.tables):
        hidden_label = " [Oculta]" if table_obj.is_hidden else ""
        _add_section_heading(doc, f"{table_obj.name}{hidden_label}", level=2)

        if table_obj.description:
            desc_p = doc.add_paragraph(table_obj.description)
            desc_p.runs[0].italic = True
            desc_p.runs[0].font.size = Pt(9)

        info_p = doc.add_paragraph()
        info_run = info_p.add_run(f"Tipo de fonte: {table_obj.source_type or 'Desconhecido'}  ·  "
                                   f"{len(table_obj.columns)} colunas  ·  {len(table_obj.measures)} medidas")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        if table_obj.columns:
            col_table = doc.add_table(rows=1, cols=4)
            col_table.style = "Table Grid"
            _add_header_row(col_table, ["Coluna", "Tipo de Dado", "Calculada?", "Descrição"])
            for i, col in enumerate(table_obj.columns):
                row = col_table.add_row()
                row.cells[0].text = col.name
                row.cells[1].text = col.data_type
                row.cells[2].text = "Sim" if col.is_calculated else "Não"
                row.cells[3].text = col.description or "—"
                _style_data_row(row, even=(i % 2 == 0))

        if table_obj.measures:
            doc.add_paragraph()
            _add_section_heading(doc, "Medidas DAX", level=3)
            for i, meas in enumerate(table_obj.measures):
                meas_p = doc.add_paragraph()
                name_run = meas_p.add_run(f"{meas.name}")
                name_run.bold = True
                name_run.font.size = Pt(10)
                name_run.font.color.rgb = COLOR_PRIMARY
                if meas.description:
                    meas_p.add_run(f"  —  {meas.description}").font.size = Pt(9)
                if meas.expression:
                    expr_p = doc.add_paragraph(meas.expression)
                    expr_p.style = doc.styles["No Spacing"]
                    for run in expr_p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x22, 0x22, 0x55)

        doc.add_paragraph()

    # ── Relacionamentos ────────────────────────────────────────────────────────
    _add_section_heading(doc, f"{section_offset}. Relacionamentos")
    section_offset += 1

    if model.relationships:
        rel_table = doc.add_table(rows=1, cols=5)
        rel_table.style = "Table Grid"
        _add_header_row(rel_table, ["De (Tabela)", "Coluna", "Para (Tabela)", "Coluna", "Cardinalidade"])
        for i, rel in enumerate(model.relationships):
            row = rel_table.add_row()
            row.cells[0].text = rel.from_table
            row.cells[1].text = rel.from_column
            row.cells[2].text = rel.to_table
            row.cells[3].text = rel.to_column
            row.cells[4].text = rel.cardinality
            _style_data_row(row, even=(i % 2 == 0))
    else:
        doc.add_paragraph("Nenhum relacionamento identificado.")

    doc.add_paragraph()

    # ── Páginas do relatório ───────────────────────────────────────────────────
    if model.report_pages:
        _add_section_heading(doc, f"{section_offset}. Páginas do Relatório")
        section_offset += 1
        for i, page in enumerate(model.report_pages, 1):
            doc.add_paragraph(f"{i}. {page}", style="List Number")
        doc.add_paragraph()

    # ── Avisos de extração ─────────────────────────────────────────────────────
    if model.extraction_warnings:
        _add_section_heading(doc, "Avisos de Extração")
        for w in model.extraction_warnings:
            p = doc.add_paragraph(f"⚠️  {w}")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0xAA, 0x55, 0x00)

    # ── Salva ──────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
